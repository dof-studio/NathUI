# visitor.py
#
# Nath UI Project
# DOF Studio/Nathmath all rights reserved
# Open sourced under Apache 2.0 License

# Backend #####################################################################

import os
import json
import csv
import chardet
import pandas as pd
from typing import Optional, Dict, List, Any
from docx import Document # pip install python-docx
from abc import ABC, abstractmethod
import pdfplumber # pip install pdfplumber
from urllib.parse import urlparse, unquote

# if you want to test, import this
from mkdown_renderer import go_renderer

class FileReaderError(Exception):
    """
    ERROR: File reader basic exception class
    """

class UnsupportedFileFormatError(FileReaderError):
    """
    ERROR: Unsupported file format exception
    """

class FileReadError(FileReaderError):
    """
    ERROR: File read failure exception
    """

# Base class: file visitor
class Base_FileVisitor(ABC):
    """
    ERROR: File reader abstract base class
    """
    
    # All extensions supported
    SUPPORTED_EXTENSIONS: Dict[str, str] = {
        '.txt': 'text',
        '.json': 'json',
        '.csv': 'csv',
        '.xls': 'excel',
        '.xlsx': 'excel',
        '.docx' : 'document',
        '.pdf': 'pdf'
    }

    def __init__(self, file_path: str):
        self._validate_file(file_path)
        self.file_path = file_path
        self.extension = self._get_extension().lower()
        self.__author__ = "DOF-Studio/NathMath@bilibili"
        self.__license__ = "Apache License Version 2.0"

    def _validate_file(self, file_path: str):
        """
        Verify that the file exists and is readable
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File NOT exists: {file_path}")
        if not os.access(file_path, os.R_OK):
            raise PermissionError(f"File Unreadable: {file_path}")
            
    def validate_file(self, file_path: str) -> bool:
        """
        Verify that the file exists and is readable, returns bool
        """
        try:
            self._validate_file(file_path)
            return True
        except:
            return False

    def _get_extension(self) -> str:
        """
        Get the file extension
        """
        _, ext = os.path.splitext(self.file_path)
        return ext

    @abstractmethod
    def read(self, as_markdown: bool = False) -> str:
        """
        Read file contents
        """
        pass

    @staticmethod
    def _to_markdown_table(data: list) -> str:
        """
        Convert a two-dimensional array to a Markdown table
        Suitable for csv or excel tables
        """
        if not data:
            return ""
        
        header = "| " + " | ".join(data[0]) + " |"
        separator = "| " + " | ".join(["---"] * len(data[0])) + " |"
        body = "\n".join(
            ["| " + " | ".join(map(str, row)) + " |" for row in data[1:]]
        )
        return "\n".join([header, separator, body])

# Deriv class: txt, plain text
class Deriv_TextFileReader(Base_FileVisitor):
    """
    Plain Text file reader (supports .txt)
    """
    
    def read(self, as_markdown: bool = False) -> str:
        encoding = self._detect_encoding()
        try:
            with open(self.file_path, "r", encoding=encoding) as f:
                content = f.read()
        except Exception as e:
            raise FileReadError(f"Failed to read the file: {str(e)}")

        return self._format_content(content, as_markdown)

    def _detect_encoding(self) -> str:
        """
        Automatically detect file encoding
        """
        try:
            with open(self.file_path, "rb") as f:
                raw = f.read(4000)
                result = chardet.detect(raw)
                return result["encoding"] or "utf-8"
        except Exception as e:
            return "utf-8"

    @staticmethod
    def _format_content(content: str, as_markdown: bool) -> str:
        """
        Format text content
        """
        return content.replace("\n", "  \n") if as_markdown else content

# Deriv class: csv 2D table
class Deriv_CsvFileReader(Base_FileVisitor):
    """
    CSV file reader
    """
    
    def read(self, as_markdown: bool = False) -> str:
        try:
            with open(self.file_path, "r", encoding="utf-8") as f:
                reader = csv.reader(f)
                data = list(reader)
        except Exception as e:
            raise FileReadError(f"Failed to read a csv file: {str(e)}")

        return self._to_markdown_table(data) if as_markdown else "\n".join([",".join(row) for row in data])

# Deriv class: json file
class Deriv_JsonFileReader(Base_FileVisitor):
    """
    JSON file reader
    """
    
    def read(self, as_markdown: bool = False) -> str:
        try:
            with open(self.file_path, "r", encoding="utf-8") as f:
                content = json.load(f)
                formatted = json.dumps(content, indent=2, ensure_ascii=False)
        except Exception as e:
            raise FileReadError(f"Failed to read a json file: {str(e)}")

        return f"```json\n{formatted}\n```" if as_markdown else formatted

# Deriv class: xls, xlsx excel spreadsheets
class Deriv_ExcelFileReader(Base_FileVisitor):
    """
    Excel spreadsheet reader (supports both .xls, .xlsx)
    """
    
    def read(self, as_markdown: bool = False) -> str:
        try:
            dfs = pd.read_excel(self.file_path, sheet_name=None)
            output = []
            for sheet_name, df in dfs.items():
                if as_markdown:
                    output.append(f"# {sheet_name}")
                    output.append(df.to_markdown(index=False))
                else:
                    output.append(f"Sheet: {sheet_name}")
                    output.append(df.to_string(index=False))
            return "\n\n".join(output)
        # By NathMath@bilibili
        except Exception as e:
            raise FileReadError(f"Failed to read the excel spreadsheet: {str(e)}")

# Deriv class: docx world document
class Deriv_WordFileReader(Base_FileVisitor):
    """
    Word file reader (only supports .docx)
    """
    
    # API 
    def read(self, as_markdown: bool = False) -> str:
        if self.extension == '.docx':
            return self._read_docx(as_markdown)
        else:
            raise UnsupportedFileFormatError(f"Unsupported format: {self.extension}")

    def _read_docx(self, as_markdown: bool) -> str:
        """
        Process .docx files
        """
        try:
            doc = Document(self.file_path)
            content = []
            
            # Process paragraphs
            for para in doc.paragraphs:
                para_text = self._process_paragraph(para, as_markdown)
                if para_text:
                    content.append(para_text)
            
            # Process tables
            for table in doc.tables:
                table_text = self._process_table(table, as_markdown)
                if table_text:
                    content.append(table_text)
            
            return '\n\n'.join(content)
        except Exception as e:
            raise FileReadError(f"Failed to read the word document: {str(e)}")

    def _process_paragraph(self, para, as_markdown: bool) -> str:
        """
        Process a single paragraph
        """
        text = self._get_style_text(para, as_markdown)
        if not text.strip():
            return ""
        
        # Markdown endline
        return text.replace('\n', '  \n') if as_markdown else text

    def _get_style_text(self, para, as_markdown: bool) -> str:
        """
        Get styled text
        """
        # Process headings
        if para.style.name.startswith('Heading'):
            level = int(para.style.name.split()[-1])
            return f"{'#' * level} {para.text.strip()}"
        
        # Process text styles (markdown)
        text_parts = []
        for run in para.runs:
            text = run.text
            if not text.strip():
                continue
            
            if as_markdown:
                if run.bold:
                    text = f"**{text}**"
                if run.italic:
                    text = f"*{text}*"
                if run.underline:
                    text = f"__{text}__"
            text_parts.append(text)
        
        return ''.join(text_parts)

    def _process_table(self, table, as_markdown: bool) -> str:
        """
        Process tables
        """
        if as_markdown:
            return self._convert_table_to_md(table)
        return self._convert_table_to_text(table)

    def _convert_table_to_md(self, table) -> str:
        """
        Convert tables to markdown
        """
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        
        if not rows:
            return ""
        
        # Add a header separator line
        separator = ['---'] * len(rows[0])
        rows.insert(1, separator)
        
        return self._to_markdown_table(rows)

    def _convert_table_to_text(self, table) -> str:
        """Convert table to plain text format"""
        return '\n'.join(
            ' | '.join(cell.text for cell in row.cells)
            for row in table.rows
        )

    def _format_doc_to_md(self, text: str) -> str:
        """
        Convert .doc text to Markdown format
        """
        return text.replace('\n', '  \n')

# Deriv class: pdf document
class Deriv_PdfFileReader(Base_FileVisitor):
    """
    PDF file reader
    """
    
    def read(self, as_markdown: bool = False) -> str:
        try:
            text = []
            with pdfplumber.open(self.file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text.append(page_text.strip())
            content = "\n\n".join(text)
        except Exception as e:
            raise FileReadError(f"Failed to read a pdf document: {str(e)}")

        return content.replace("\n", "  \n") if as_markdown else content

# Generic File visitor class (supports all extensions)
class Generic_FileVisitor:
    """
    Generic File reader factory class
    """
    
    @staticmethod
    def create_reader(file_path: str) -> Base_FileVisitor:
        """
        Create a corresponding reader instance based on the file extension
        """
        ext = os.path.splitext(file_path)[1].lower()
        readers = {
            ".txt": Deriv_TextFileReader,
            ".csv": Deriv_CsvFileReader,
            ".json": Deriv_JsonFileReader,
            ".xls": Deriv_ExcelFileReader,
            ".xlsx": Deriv_ExcelFileReader,
            ".docx": Deriv_WordFileReader,
            ".pdf": Deriv_PdfFileReader,
        }
        
        if ext not in readers:
            # Regard as txt
            return readers[".txt"](file_path)
        
        return readers[ext](file_path)

# API: File visitor (noexcept)
def file_visitor(file_path: str, as_markdown: bool = False, noexcept: bool = True) -> str:
    """
    Unified file reading interface
    """
    try:
        reader = Generic_FileVisitor.create_reader(file_path)
        return reader.read(as_markdown=as_markdown)
    except Exception as e:
        if noexcept:
            return "" # Read nothing
        else:
            raise FileReaderError(str(e))
            
# API: Is File (Url, Folder, ...)
def is_file(anything: str) -> int:
    """
    See whether a string contains a file or an url
    Return 0(nonexist file), 1(existing file), 2(url), 3(existing folder), -1(others)
    """
    s = anything.strip()
    parsed = urlparse(s)
    scheme = parsed.scheme.lower()
    
    # Constant defines
    _nfile   =  0
    _file    =  1
    _url     =  2
    _folder  =  3
    _other   = -1

    if scheme:
        if scheme == 'file':
            # For file URLs, convert the URL path to a local file path.
            path = unquote(parsed.path)
            if os.name == 'nt':
                # On Windows, remove a leading slash if it precedes a drive letter.
                if path.startswith('/') and len(path) > 1 and path[2] == ':':
                    path = path.lstrip('/')
            if os.path.exists(path):
                return _folder if os.path.isdir(path) else _file
            else:
                return _nfile
        elif scheme in ('http', 'https', 'ftp', 'udp'):
            return _url
        else:
            # Unknown scheme; treat as local path.
            if os.path.exists(s):
                return _folder if os.path.isdir(s) else _file
            else:
                return _other
    else:
        # No scheme provided; assume it's a local file path.
        if os.path.exists(s):
            return _folder if os.path.isdir(s) else _file
        else:
            return _other


# Test
if __name__ == "__main__":
    # 读取文本文件
    # print(file_visitor("__testsample__/text.txt", as_markdown=False))
    
    # 读取CSV并转换为Markdown表格
    # print(file_visitor("__testsample__/marketcap.csv", as_markdown=True))
    # go_renderer(file_visitor("__testsample__/marketcap.csv", as_markdown=True))
    
    # 读取Excel文件
    # print(file_visitor("__testsample__/normsinv.xlsx",as_markdown=True))
    # go_renderer(file_visitor("__testsample__/normsinv.xlsx", as_markdown=True))

    # 读取PDF文件
    # print(file_visitor("__testsample__/montecarlo.pdf", as_markdown=False))
    # go_renderer(file_visitor("__testsample__/montecarlo.pdf", as_markdown=True))

    # 读取.docx文件
    # print(file_visitor("__testsample__/exampaper.docx", as_markdown=True))
    # go_renderer(file_visitor("__testsample__/exampaper.docx", as_markdown=True))


    pass